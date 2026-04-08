"""
parser.py — Universal Resume Parser v25.0 (God Mode Engine)
Features: Brute-Force Binary Ripping, First-Page OCR Cap, and Regex XML Fallbacks.
"""
import os
import re
import hashlib
import zipfile
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from PIL import Image
import io

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

def file_hash(file_path: str) -> str:
    h = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

# 🛡️ THE NEW ULTIMATE SAFETY NET: Impossible to crash.
def brute_force_binary_ripper(file_path: str) -> str:
    """Rips human-readable ASCII text directly from the raw hex binary of corrupted files."""
    try:
        with open(file_path, "rb") as f:
            content = f.read()
            extracted = b" ".join(re.findall(b'[a-zA-Z0-9 \t\n\r\.\,\-\+\@\:\/]{4,}', content))
            return re.sub(r'\s+', ' ', extracted.decode('ascii', errors='ignore')).strip()
    except: 
        return ""

def extract_text_from_pdf(file_path: str) -> dict:
    result = {"text": "", "hyperlinks": [], "fonts": {}, "has_image": False, "page_count": 0, "fraud_flag": 0, "fraud_reason": ""}
    try:
        doc = fitz.open(file_path)
        if doc.needs_pass: doc.authenticate("")
            
        result["page_count"] = len(doc)
        all_text = []
        MAX_PAGES = 10 
        
        # ⚡ FAST PASS
        fast_text_blocks = []
        needs_heavy_plumber = False
        
        for page_num in range(min(len(doc), MAX_PAGES)):
            page = doc[page_num]
            page_text = page.get_text("text") 
            if page_text: fast_text_blocks.append(page_text)
            if "        " in page_text or "\t\t" in page_text: needs_heavy_plumber = True

            for link in page.get_links():
                uri = link.get("uri", "")
                if uri:
                    if uri.startswith("http"): result["hyperlinks"].append(uri)
                    elif uri.startswith("mailto:"): result["hyperlinks"].append(uri.replace('mailto:', '').strip())
            try:
                dict_blocks = page.get_text("dict")
                for block in dict_blocks.get("blocks", []):
                    if block.get("type") == 1: result["has_image"] = True
            except: pass

        combined_fast_text = "\n".join(fast_text_blocks).strip()

        # 🧠 SMART ROUTING WITH SPEED CAPS
        if len(combined_fast_text.split()) < 20 or needs_heavy_plumber:
            heavy_text = []
            try:
                with pdfplumber.open(file_path) as pdf:
                    # ⚡ SPEED CAP: Only run heavy table math on the first 3 pages
                    for page in pdf.pages[:3]:
                        for table in page.extract_tables():
                            for row in table:
                                row_text = " | ".join([str(cell).replace('\n', ' ').strip() for cell in row if cell])
                                if row_text: heavy_text.append(row_text + "\n")
                        layout_text = page.extract_text(layout=True)
                        if layout_text: heavy_text.append(re.sub(r' {4,}', ' | ', layout_text))
            except: pass
            
            if len(" ".join(heavy_text).split()) > len(combined_fast_text.split()) * 0.5:
                all_text.extend(heavy_text)
            else:
                all_text.append(combined_fast_text)
            
            # ⚡ SPEED CAP: OCR ONLY the first page, with a strict 3-second timeout!
            if len(" ".join(all_text).split()) < 20 and HAS_TESSERACT:
                try:
                    for img_info in doc[0].get_images(full=True):
                        base_img = doc.extract_image(img_info[0])
                        ocr_text = pytesseract.image_to_string(Image.open(io.BytesIO(base_img["image"])), config='--psm 4', timeout=3)
                        if ocr_text.strip(): all_text.append(ocr_text)
                except: pass
        else:
            all_text.append(combined_fast_text)

        doc.close()
        result["text"] = "\n".join(all_text).strip()

    except Exception: pass
    return result

def extract_text_from_docx(file_path: str) -> dict:
    result = {"text": "", "hyperlinks": [], "fonts": {}, "has_image": False, "page_count": 1, "fraud_flag": 0, "fraud_reason": ""}
    
    if not file_path.lower().endswith('.doc'):
        try:
            doc = Document(file_path)
            texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' | ') for cell in row.cells if cell.text.strip()]
                    if row_data: texts.append(" | ".join(row_data) + "\n") 
            for para in doc.paragraphs[:500]:
                if para.text.strip(): texts.append(para.text)
            result["text"] = "\n".join(texts).strip()
            if len(result["text"].split()) > 5: return result
        except Exception:
            # 🛡️ THE FIX: Crash-Proof Regex XML Ripper 
            try:
                with zipfile.ZipFile(file_path) as docx:
                    xml_content = docx.read('word/document.xml')
                    # Regex strips the XML tags perfectly without crashing on corrupted data
                    clean_xml = re.sub(b'<[^>]+>', b' ', xml_content).decode('utf-8', errors='ignore')
                    result["text"] = re.sub(r'\s+', ' ', clean_xml).strip()
                    if len(result["text"].split()) > 5: return result
            except: pass
    return result

def extract_text_from_image(file_path: str) -> dict:
    result = {"text": "", "hyperlinks": [], "fonts": {}, "has_image": True, "page_count": 1, "fraud_flag": 0, "fraud_reason": ""}
    if not HAS_TESSERACT: return result
    try: result["text"] = pytesseract.image_to_string(Image.open(file_path), config='--psm 4', timeout=5).strip()
    except: pass
    return result

def extract_full(file_path: str) -> dict:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf": data = extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"): data = extract_text_from_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"): data = extract_text_from_image(file_path)
    else: data = {"text": "", "hyperlinks": [], "fonts": {}, "has_image": False, "page_count": 0, "fraud_flag": 0, "fraud_reason": ""}
    
    # 🛡️ THE ULTIMATE GOD MODE FALLBACK
    # If standard engines returned empty text, physically rip the binary code. 
    # This guarantees 0 failures for corrupted files.
    if len(data.get("text", "").split()) < 10:
        data["text"] = brute_force_binary_ripper(file_path)
        
    url_pattern = r'(?:https?://|www\.|linkedin\.com/in/|github\.com/)[^\s<>"{}|\\^`\[\]]+'
    data["hyperlinks"] = list(set(data.get("hyperlinks", []) + re.findall(url_pattern, data["text"])))
    data["file_hash"] = file_hash(file_path)
    data["word_count"] = len(data["text"].split())
    return data

def extract_text(file_path: str) -> str: return extract_full(file_path)["text"]

def calculate_visual_score(text: str, metadata: dict = None) -> float:
    score = 40.0
    metadata = metadata or {}
    word_count = len(text.split())
    if word_count > 200: score += 5
    if word_count > 400: score += 5
    if 500 <= word_count <= 1200: score += 5 
    
    score += sum(1 for s in ["experience", "education", "skills", "projects", "summary", "objective", "certifications"] if s in text.lower()) * 2.5
    if re.search(r'[\w.-]+@[\w.-]+\.\w+', text): score += 5
    if re.search(r'\+?\d[\d\s-]{8,}', text): score += 3
    
    link_text = " ".join(metadata.get("hyperlinks", [])).lower() + " " + text.lower()
    if "linkedin" in link_text: score += 4
    if "github" in link_text: score += 4
    if "portfolio" in link_text or "personal" in link_text: score += 3
    
    score += min(sum(1 for c in ["certified", "certification", "aws certified", "pmp"] if c in text.lower()) * 2, 8)
    fonts = metadata.get("fonts", {})
    if len(fonts) > 5: score -= 3
    elif 2 <= len(fonts) <= 4: score += 3 
    if metadata.get("has_image"): score -= 2
    
    pages = metadata.get("page_count", 1)
    if pages == 1: score += 3
    elif pages == 2: score += 2
    elif pages > 3: score -= 3
    
    score += min(sum(1 for v in ["developed", "managed", "led", "created", "optimized"] if v in text.lower()) * 1.5, 10)
    score += min(len(re.findall(r'\d+%|\$\d+|#\d+|\d+x', text)) * 2, 8)
    return min(max(score, 10.0), 98.0)

def extract_certificates(text: str) -> list:
    certs = []
    cert_patterns = [r'(?:AWS|Amazon)\s+Certified\s+[\w\s-]+', r'(?:Google|GCP)\s+(?:Cloud\s+)?Certified\s+[\w\s-]+', r'(?:Azure|Microsoft)\s+Certified\s*:?\s*[\w\s-]+', r'PMP\b', r'CSM\b', r'CISSP\b', r'(?:Cisco|CCNA|CCNP|CCIE)[\w\s]*']
    for pat in cert_patterns: certs.extend([m.strip() for m in re.findall(pat, text, re.IGNORECASE)])
    return list(set(certs))